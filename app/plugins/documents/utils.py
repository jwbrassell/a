"""Utility functions for document management."""
import os
from datetime import datetime, timedelta
import pdfkit
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from app import db
from app.plugins.documents.models import Document, DocumentCache

def get_or_create_cache(document, format='html', force_refresh=False):
    """Get cached version of document or create new cache."""
    # Check for existing non-expired cache
    cache = DocumentCache.query.filter_by(
        document_id=document.id,
        format=format
    ).filter(
        DocumentCache.expires_at > datetime.utcnow()
    ).first()

    if cache and not force_refresh:
        cache.access_count += 1
        db.session.commit()
        return cache.content

    # Create new cache
    if format == 'html':
        content = document.content
    elif format == 'pdf':
        content = create_pdf(document)
    elif format == 'doc':
        content = create_doc(document)
    else:
        raise ValueError(f"Unsupported format: {format}")

    # Delete old caches for this format
    DocumentCache.query.filter_by(
        document_id=document.id,
        format=format
    ).delete()

    # Create new cache entry
    cache = DocumentCache(
        document_id=document.id,
        content=content,
        format=format,
        expires_at=datetime.utcnow() + timedelta(days=1)  # Cache for 24 hours
    )
    db.session.add(cache)
    db.session.commit()

    return content

def create_pdf(document):
    """Convert document HTML to PDF."""
    try:
        # Configure PDF options
        options = {
            'page-size': 'Letter',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8"
        }
        
        # Create PDF from HTML content
        pdf_content = pdfkit.from_string(document.content, False, options=options)
        return pdf_content
        
    except Exception as e:
        print(f"Error creating PDF: {str(e)}")
        return None

def create_doc(document):
    """Convert document HTML to DOCX format."""
    try:
        # Create new Word document
        doc = DocxDocument()
        doc.add_heading(document.title, 0)
        
        # Parse HTML content
        soup = BeautifulSoup(document.content, 'html.parser')
        
        # Convert HTML to Word paragraphs
        for p in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            if p.name.startswith('h'):
                level = int(p.name[1])
                doc.add_heading(p.get_text(), level)
            else:
                doc.add_paragraph(p.get_text())
        
        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error creating DOC: {str(e)}")
        return None

def bulk_categorize(document_ids, category_id):
    """Bulk update document categories."""
    try:
        Document.query.filter(Document.id.in_(document_ids)).update(
            {Document.category_id: category_id},
            synchronize_session=False
        )
        db.session.commit()
        return True
    except Exception as e:
        db.session.rollback()
        print(f"Error bulk categorizing documents: {str(e)}")
        return False

def bulk_delete(document_ids):
    """Bulk delete documents."""
    try:
        # Delete associated caches first
        DocumentCache.query.filter(
            DocumentCache.document_id.in_(document_ids)
        ).delete(synchronize_session=False)
        
        # Delete documents
        Document.query.filter(Document.id.in_(document_ids)).delete(
            synchronize_session=False
        )
        
        db.session.commit()
        return True
    except Exception as e:
        db.session.rollback()
        print(f"Error bulk deleting documents: {str(e)}")
        return False
